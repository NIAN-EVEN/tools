import itchat
import copy, os
from itchat.content import *
from docx import Document
from docx.shared import Inches

# 设置检测的群聊名称和检测关键字
# groupName = "第十七期北辰青年领袖营✨"
groupName = "北十七后厨"
keyWord = "日新笔记"

@itchat.msg_register(itchat.content.TEXT, isGroupChat=True)
def store_msg(msg):
    # 接收消息
    path = "日新笔记/"
    lev = 1
    # 获取目标群的id
    group = itchat.search_chatrooms(name=groupName)
    if len(group) != 1:
        print("There is a lot of groups named ", groupName)
        print("Please readjust group name then rerun the program")
        exit(-1)

    groupId = group[0]["UserName"]
    # 获取目标群成员信息
    groupMbr = {}
    for mbr in group[0]["MemberList"]:
        print(mbr["UserName"], ' ', mbr["NickName"])
        groupMbr[mbr["UserName"]] = mbr["NickName"]
    print("="*50)
    # 判断是否来自目标群聊
    if msg["FromUserName"] == groupId:
        # 如果是来自于希望的群聊
        # 判断发送的消息中是否含有希望的关键字
        if "日新笔记" in msg["Content"]:
            # 如果包含希望的关键字
            # 获取发送人信息
            print(msg["ActualNickName"])
            sender = groupMbr[msg["ActualUserName"]]
            # 判断文件夹内是否有该发送人的文档
            if os.path.exists(path+sender+".docx"):
                # 如果有
                # 追加文本到文档内
                document = Document(path+sender+".docx")
                document.add_paragraph(msg["Content"])
                document.save(path + sender + ".docx")
            else:
                # 如果没有
                # 创建docx文档
                document = Document()
                # 追加文本到文档内
                document.add_paragraph(msg["Content"])
                document.add_paragraph("\n")
                # 保存文本
                document.save(path + sender + ".docx")
        else:
            print("not desired msg")
            print("="*50)
            # 如果不包含希望的关键字：重新监控消息
            pass

    else:
        print("not desired group")
        print("="*50)
        # 如果不是来自于希望的群聊：重新监控消息
        pass


def transMsgContent(contents):
    content = copy.deepcopy(contents)
    for i, chr in enumerate(content):
        if chr in range(0x1f600, 0x1f650):
            contents[i] == ' '
    return content

if __name__ == "__main__":
    # 登录
    itchat.auto_login(hotReload=True)
    itchat.run()
    itchat.get_msg()
